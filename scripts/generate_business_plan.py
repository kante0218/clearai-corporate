"""clear AI株式会社 事業計画書 生成スクリプト"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_default_font(doc, font_name="游明朝"):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "游ゴシック"
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "游ゴシック")
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return h


def add_para(doc, text, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    r.font.name = "游明朝"
    rpr = r._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "游明朝")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.name = "游明朝"
        rpr = r._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "游明朝")
    return p


def add_table(doc, headers, rows, header_bg="1E3A8A", header_color=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = header_color
        r.font.name = "游ゴシック"
        rpr = r._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "游ゴシック")
        set_cell_bg(hdr[i], header_bg)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "游明朝"
            rpr = r._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "游明朝")
    return table


def main():
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── 表紙 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("\n\n\nclear AI 株式会社\n事業計画書")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    r.font.name = "游ゴシック"
    rpr = r._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "游ゴシック")

    add_para(doc, "\n学業と両立した、現実的な5カ年成長計画（2026–2030）", bold=True, size=13, align="center", color=RGBColor(0x25, 0x63, 0xEB))
    add_para(doc, "\n\n\n\n\n作成日：2026年4月 ／ Version 1.0", size=10, align="center", color=RGBColor(0x6B, 0x72, 0x80))
    add_para(doc, "clear AI 株式会社　代表取締役　髙橋 敢輝", size=10, align="center", color=RGBColor(0x6B, 0x72, 0x80))

    doc.add_page_break()

    # ── 0. エグゼクティブサマリー ──
    add_heading(doc, "0. エグゼクティブサマリー", 1)
    add_para(doc, "clear AI株式会社は、代表が筑波大学での学業と両立しながら経営する、AIコンサルティング＋農業×エンジニアリングの二本柱スタートアップです。本計画書は、従前に掲げていた「2030年売上1000億円」という目標を現実的なラインへ再設定し、学業との両立を前提とした5カ年の段階的成長計画を示します。")
    add_para(doc, "")
    add_para(doc, "【見直し後の目標】", bold=True)
    add_bullet(doc, "2030年 売上高：10億円（従前の1/100規模に現実化、ただし粗利率50%以上の高収益構造を志向）")
    add_bullet(doc, "2030年 AI顧問契約：40社／MRR 2,500万円（主力リカーリング収益源）")
    add_bullet(doc, "2030年 累計支援企業数：120社（うち継続契約60社）")
    add_bullet(doc, "2028年 農業事業：黒字転換（単月黒字を2027年中、通期黒字を2028年度に）")
    add_bullet(doc, "代表の稼働：在学中は週30時間上限／2028年度（卒業後）からフルタイム経営へ移行")

    doc.add_page_break()

    # ── 1. ビジョンの再定義 ──
    add_heading(doc, "1. ビジョンの再定義", 1)

    add_heading(doc, "1.1 なぜ「1000億円 / 10,000社」を見直すのか", 2)
    add_para(doc, "創業時に掲げた「2030年売上1000億円・支援企業10,000社」という目標は、日本のAIコンサル業界の実態（アクセンチュア、野村総研クラスの規模）と比較しても非現実的であり、学業と両立する経営形態とは整合しません。また、短期で無理な拡大を図ることは、以下のリスクを伴います。")
    add_bullet(doc, "低品質な納品による顧客離反と、4つの信条（誠実・伴走・翻訳・長期視点）との乖離")
    add_bullet(doc, "代表の学業成績悪化、研究機会の喪失")
    add_bullet(doc, "資金調達への過度な依存と、創業者持株比率の過度な希薄化")
    add_para(doc, "本計画では、「数字の壮大さ」ではなく「顧客1社あたりの成功事例の深さ」を競争優位の源泉と再定義します。")

    add_heading(doc, "1.2 新ビジョン・ステートメント", 2)
    add_para(doc, "「AIを、日本の現場へ届ける。」というコア・メッセージは維持しつつ、以下を新しい定量ビジョンとします。", bold=True)
    add_para(doc, "")
    add_table(
        doc,
        ["時間軸", "ビジョン", "定量指標"],
        [
            ["2030年", "茨城発・AI×地域産業の専門ブティック", "売上10億円／粗利5億円／社員15名"],
            ["2035年", "関東圏で「AI×一次産業」と言えばclear AI", "売上50億円／社員50名"],
            ["2040年〜", "食料安全保障インフラを担う中堅テック企業", "売上200億円規模（IPO視野）"],
        ],
    )

    add_heading(doc, "1.3 1000億円という数字の扱い", 2)
    add_para(doc, "「1000億円」は、2040年代の超長期の野心的マイルストーンとして社内のみで保持し、対外公表資料（コーポレートサイト等）からは数値目標を「2030年：売上10億円／支援企業120社／47都道府県中20都道府県」に差し替えます。過度な数値目標が顧客・採用候補者に与える不信感を回避し、実績の積み上げで信頼を勝ち取る戦略へ転換します。")

    doc.add_page_break()

    # ── 2. 5カ年ロードマップ ──
    add_heading(doc, "2. 5カ年ロードマップ（学業との両立前提）", 1)
    add_para(doc, "代表の筑波大学在学期間（2026–2028年度）は「基盤構築フェーズ」として、規模よりも事例と仕組みづくりを優先します。卒業後の2029年度から「拡大フェーズ」に移行し、採用と営業投資を本格化します。")
    add_para(doc, "")

    add_table(
        doc,
        ["年度", "フェーズ", "代表稼働", "売上目標", "主要マイルストーン"],
        [
            ["2026年度", "創業・PMF検証", "週20h（学業優先）", "3,000万円", "AI顧問契約3社、農業EC案件2件納品、社員0→2名"],
            ["2027年度", "事例蓄積", "週25h", "8,000万円", "AI顧問10社、農業事業単月黒字化、社員2→5名"],
            ["2028年度", "仕組み化", "週30h→卒業後フル", "2.0億円", "AI顧問20社、農業通期黒字、初の新卒採用"],
            ["2029年度", "拡大初年度", "フルタイム", "4.5億円", "関東圏営業拠点、社員8→15名、シリーズA検討"],
            ["2030年度", "地域No.1確立", "フルタイム", "10.0億円", "支援120社、粗利率50%以上、2拠点体制"],
        ],
    )
    add_para(doc, "")
    add_para(doc, "【売上の内訳想定（2030年度／単位：百万円）】", bold=True)
    add_table(
        doc,
        ["事業", "売上", "粗利率", "備考"],
        [
            ["AI顧問サービス（月額）", "300", "70%", "40社 / 平均月62万円 / MRR2,500万円"],
            ["AI導入プロジェクト（受託）", "460", "40%", "年12件 × 平均3,800万円（顧問先からが75%）"],
            ["農業×エンジニアリング事業", "200", "35%", "EC構築＋運用受託＋自社EC販売"],
            ["研修・ライセンス", "40", "65%", "法人向けAIリテラシー研修"],
            ["合計", "1,000", "約50%", ""],
        ],
    )

    doc.add_page_break()

    # ── 3. 事業別戦略 ──
    add_heading(doc, "3. 事業別戦略", 1)

    add_heading(doc, "3.1 AIコンサルティング事業（主力：AI顧問サービス）", 2)
    add_para(doc, "既に「日本アセット戦略機構」にAI顧問として導入実績があり、月額リカーリング収益の核として最優先で拡大します。顧問契約はキャッシュフローが安定し、代表の稼働も月数回の定例＋チャットで収まるため、学業と最も両立しやすいビジネスモデルです。ここを事業の「幹」と位置付けます。")

    add_para(doc, "")
    add_para(doc, "【顧問契約の3プラン体系】", bold=True)
    add_table(
        doc,
        ["プラン", "月額", "提供内容", "想定顧客", "代表稼働／社"],
        [
            ["Starter", "30万円", "月2回の定例オンライン + チャット相談", "売上10–50億円の中小企業", "月3–4h"],
            ["Standard", "80万円", "業務棚卸 + プロンプト設計 + 簡易実装", "売上50–300億円の中堅企業", "月6–8h"],
            ["Enterprise", "150–200万円", "部門横断の変革伴走 + 経営会議参画", "売上300億円以上 / 金融・製造", "月10–12h"],
        ],
    )

    add_para(doc, "")
    add_para(doc, "【顧問契約の拡大目標（最重点KPI）】", bold=True)
    add_table(
        doc,
        ["年度", "顧問社数", "内訳（S/St/E）", "MRR", "年換算売上"],
        [
            ["2026", "3社", "2/1/0", "90万円", "1,080万円"],
            ["2027", "10社", "6/3/1", "430万円", "5,160万円"],
            ["2028", "20社", "10/8/2", "1,040万円", "1.25億円"],
            ["2029", "30社", "14/12/4", "1,720万円", "2.06億円"],
            ["2030", "40社", "18/16/6", "2,500万円", "3.00億円"],
        ],
    )
    add_para(doc, "※顧問契約は解約率月1%以下を前提。粗利率70%で、2030年度は顧問単独で粗利2.1億円を創出します。", size=9, color=RGBColor(0x6B, 0x72, 0x80))

    add_para(doc, "")
    add_para(doc, "【顧問契約を増やすための5つの打ち手】", bold=True)
    add_bullet(doc, "① 既存顧客（日本アセット戦略機構 等）を起点にした紹介営業：成功事例を匿名化して紹介用資料化し、紹介1件あたりの成約率を30%→50%に引き上げる")
    add_bullet(doc, "② 地場金融機関・商工会議所との提携：茨城県内の常陽銀行・筑波銀行の取引先向けに「AI顧問紹介プログラム」を持ち込み、銀行から月1–2社の紹介ルートを構築")
    add_bullet(doc, "③ 士業ネットワーク連携：中小企業診断士・税理士・社労士と業務提携し、顧問先を相互紹介（フィー20%分配）")
    add_bullet(doc, "④ コンテンツマーケティング：代表名義で「現場で使えるAI活用事例」を週1本発信、2027年度以降Web経由で月2社の問い合わせを獲得")
    add_bullet(doc, "⑤ 業界特化パッケージ：まず「金融・不動産」「製造」「農業関連」の3業界で事例を深掘りし、業界別の顧問パッケージとして横展開（1事例→10社への展開効率）")

    add_para(doc, "")
    add_para(doc, "【顧問契約 → 実装プロジェクトへのアップセル】", bold=True)
    add_bullet(doc, "顧問先の50%以上から年1件の実装プロジェクト（平均3,000万円）を受注することを目標")
    add_bullet(doc, "顧問は「信頼関係を結ぶ継続接点」、実装は「まとまった売上を作る刈り取り接点」と役割分担")
    add_bullet(doc, "2030年度の実装案件40件中、30件（75%）は顧問先からの受注を想定")

    add_heading(doc, "3.2 農業×エンジニアリング事業", 2)
    add_para(doc, "「日本の食料安全保障」という長期ビジョンへ繋がる事業。ただし、直接の農業経営は参入せず、あくまで「農家の経営をエンジニアリングで支える」サービス業として設計することで、資本コストと在庫リスクを抑えます。")
    add_para(doc, "詳細な黒字化プランは次章に記載します。")

    doc.add_page_break()

    # ── 4. 農業事業の黒字化プラン ──
    add_heading(doc, "4. 農業×エンジニアリング事業の黒字化プラン", 1)

    add_heading(doc, "4.1 現状認識", 2)
    add_para(doc, "農業事業は社会的意義が高い一方、単価が低く、顧客（農家）のITリテラシーとキャッシュフローが事業成長のボトルネックになりがちです。この構造的課題を直視した上で、以下の3段階アプローチで確実な黒字化を目指します。")

    add_heading(doc, "4.2 黒字化の3段構造", 2)

    add_para(doc, "【Step 1】一次請けコンサル＋EC構築（2026–2027）", bold=True, color=RGBColor(0x25, 0x63, 0xEB))
    add_para(doc, "低リスク・高粗利のサービス受託で赤字を小さく抑え、農業現場のドメイン知識と信頼を獲得するフェーズ。")
    add_bullet(doc, "Shopify / BASE等を使った農家向けECサイト構築（単価40–80万円、粗利率60%）")
    add_bullet(doc, "ブランディング・商品撮影・LP制作パッケージ（単価30–50万円）")
    add_bullet(doc, "年間20件の納品で売上1,000–1,500万円、粗利600–900万円を狙う")
    add_bullet(doc, "JA・農業法人・道の駅・地銀の営業網と提携して営業コストを圧縮")

    add_para(doc, "")
    add_para(doc, "【Step 2】月額運用支援＋D2C代行（2027–2028）", bold=True, color=RGBColor(0x25, 0x63, 0xEB))
    add_para(doc, "一度作ったECを活用し、継続課金モデルに転換。単月黒字を達成する最重要フェーズ。")
    add_bullet(doc, "EC運用代行：月額5–15万円 × 30農家（AI活用で商品説明文生成・需要予測・SNS運用を自動化）")
    add_bullet(doc, "D2C代行（売上連動）：売上の15–20%を手数料として徴収。仕入れゼロのレベニューシェア型")
    add_bullet(doc, "2027年度中に月額MRR 300万円を確立、農業事業単体で単月黒字化")

    add_para(doc, "")
    add_para(doc, "【Step 3】産地共同ブランド「clear farm」プラットフォーム（2028–2030）", bold=True, color=RGBColor(0x25, 0x63, 0xEB))
    add_para(doc, "支援先農家を束ねた共同ブランドを立ち上げ、個別農家では不可能な企業向け・輸出向け大口販売を代行。これにより通期黒字化と、将来の産業インフラ化への布石を打つ。")
    add_bullet(doc, "茨城県産の高付加価値野菜・米を「clear farm」ブランドで共同出荷")
    add_bullet(doc, "法人向け（社員食堂・ふるさと納税返礼・高級スーパー）へ卸売営業")
    add_bullet(doc, "需給予測AIで廃棄ロスを削減し、個別農家比で手取り20%アップを実現")
    add_bullet(doc, "2030年度に農業事業単体で売上2億円・営業利益率15%を目指す")

    add_heading(doc, "4.3 黒字化のP/L試算（農業事業単体、単位：百万円）", 2)
    add_table(
        doc,
        ["年度", "売上", "売上原価", "販管費", "営業損益", "状態"],
        [
            ["2026", "15", "6", "18", "▲9", "赤字（投資期）"],
            ["2027", "40", "18", "25", "▲3", "単月黒字達成"],
            ["2028", "90", "45", "40", "+5", "通期黒字転換"],
            ["2029", "140", "75", "50", "+15", "安定黒字"],
            ["2030", "200", "110", "60", "+30", "スケール期"],
        ],
    )

    add_heading(doc, "4.4 黒字化を阻む3大リスクと対策", 2)
    add_bullet(doc, "リスク①：農家のITリテラシー不足 → 現地訪問＋動画マニュアル＋電話サポートを標準化し、解約率を月次1%未満に抑える")
    add_bullet(doc, "リスク②：季節変動によるキャッシュ不安定 → 月額契約を必ず12ヶ月一括 or 自動更新とし、売上連動フィーは上限設定")
    add_bullet(doc, "リスク③：代表の学業繁忙期（試験期間等）の品質低下 → 業務フロー標準化 + パートタイム運用スタッフ（農学系学生バイト）の体制を2027年度に完成")

    doc.add_page_break()

    # ── 5. 組織・人員計画 ──
    add_heading(doc, "5. 組織・人員計画（学業両立モデル）", 1)
    add_para(doc, "学業と両立するためには、代表依存を最小化する組織設計が鍵となります。2027年度中に「代表がいない週があっても事業が止まらない」体制を完成させます。")
    add_table(
        doc,
        ["年度", "社員", "業務委託", "学生インターン", "代表の主業務"],
        [
            ["2026", "2名", "3名", "1名", "営業・顧問サービス提供"],
            ["2027", "5名", "5名", "3名", "営業・採用・パートナー開拓"],
            ["2028", "8名", "6名", "5名", "経営全般（卒業後フル）"],
            ["2029", "15名", "8名", "6名", "経営全般・シリーズA"],
            ["2030", "25名", "10名", "8名", "経営全般・IPO準備初期"],
        ],
    )
    add_para(doc, "")
    add_para(doc, "【学生インターンを戦略的に活用】", bold=True)
    add_bullet(doc, "筑波大学・茨城大学の学生を農業事業のEC運用支援・データ入力で積極活用")
    add_bullet(doc, "将来の新卒採用パイプラインと位置付け、内定直結のインターン設計とする")

    doc.add_page_break()

    # ── 6. 資金計画 ──
    add_heading(doc, "6. 資金計画", 1)
    add_para(doc, "在学中はエクイティ調達を極力回避し、自己資金＋補助金＋営業キャッシュフローで回す方針。持株比率を維持しながら、黒字化後に必要最小限の調達を検討します。")
    add_table(
        doc,
        ["年度", "調達方法", "金額", "用途"],
        [
            ["2026", "自己資金＋日本政策金融公庫（創業融資）", "500万円", "初期運転資金"],
            ["2027", "IT導入補助金／ものづくり補助金", "300万円", "自社プロダクト開発"],
            ["2028", "営業キャッシュフロー中心", "ー", "人件費・採用投資"],
            ["2029", "シリーズA検討（1–2億円、評価額10億円前後）", "1.5億円", "営業拠点・エンジニア採用"],
            ["2030", "ー", "ー", "自己資金で運営、次ラウンド準備"],
        ],
    )

    doc.add_page_break()

    # ── 7. KPI ──
    add_heading(doc, "7. 月次で追う主要KPI", 1)
    add_bullet(doc, "【AI顧問】MRR（月次経常収益）、解約率、顧客単価、NPS")
    add_bullet(doc, "【農業】取引農家数、EC構築受注数、農業事業MRR、農家あたり手取り向上率")
    add_bullet(doc, "【組織】代表稼働時間（週30h上限を死守）、社員稼働時間、採用進捗")
    add_bullet(doc, "【学業】代表の単位取得数・GPA（「学業が止まらないこと」を経営指標に組み込む）")

    doc.add_page_break()

    # ── 8. 終わりに ──
    add_heading(doc, "8. 終わりに", 1)
    add_para(doc, "派手さより、10年続く仕組みを。これはclear AIの信条「長期視点」そのものです。「売上1000億円」という数字を掲げることは容易ですが、数字が独り歩きした結果、品質と信頼を失うスタートアップを私たちは数多く見てきました。")
    add_para(doc, "本計画書は、学業を投げ出さず、顧客を軽視せず、社員を使い捨てない成長をコミットするためのドキュメントです。2030年の売上10億円は控えめに見えるかもしれません。しかし、粗利率50%・継続率95%・NPS50超という「質」を伴った10億円は、無理に積み上げた100億円より遥かに強く、次の100億円、1000億円への揺るぎない踏み台になると確信しています。")
    add_para(doc, "")
    add_para(doc, "clear AI 株式会社　代表取締役　髙橋 敢輝", bold=True, align="center", color=RGBColor(0x1E, 0x3A, 0x8A))

    out = "/Users/user/Desktop/clearai-corporate/clearAI_事業計画書.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
