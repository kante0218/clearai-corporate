#!/usr/bin/env python3
"""各社長ごとに個別のWord手順書を作成"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.dirname(__file__))


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Yu Gothic'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return doc


def add_title_page(doc, name, subtitle, tagline):
    for _ in range(4):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{tagline}')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI活用 実践手順書  |  clearAI')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)
    run.font.bold = True

    doc.add_page_break()


def add_step_section(doc, step_num, title, tool, steps, tips):
    doc.add_heading(f'STEP {step_num}: {title}', level=2)

    p = doc.add_paragraph()
    run = p.add_run(f'使用ツール: {tool}')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)
    run.font.size = Pt(12)

    doc.add_paragraph('')
    doc.add_heading('手順', level=3)
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph('')
    doc.add_heading('ポイント・コツ', level=3)
    for tip in tips:
        if tip:
            p = doc.add_paragraph(tip, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph('')


def save_doc(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Saved: {path}")


# ============================================================
# 1. 南場智子
# ============================================================
def create_nanba():
    doc = setup_doc()
    add_title_page(doc,
        '南場 智子（DeNA会長）のAI活用術',
        'Perplexity × NotebookLM × Circleback で情報戦を制する',
        '「経営者がAIに興奮しているかがポイント」')

    doc.add_heading('南場流 AI活用の全体像', level=1)
    doc.add_paragraph(
        'DeNA会長の南場智子氏は、2025年に「DeNAはAIにオールイン」を宣言し、'
        '自ら生成AIのヘビーユーザーとして具体的な活用法を「7つのAI活用術」として公開しました。'
        '初めて会う人のリサーチからミーティングの議事録、投資判断まで、あらゆる場面でAIを活用しています。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        'Perplexity AIで「必読記事」を見つける',
        'Perplexity AI（perplexity.ai）',
        [
            'Perplexity AI（perplexity.ai）を開く',
            '「〇〇さんについての必読記事は何ですか？」と入力',
            '表示された記事のURLをすべてコピーする',
            'X（Twitter）の発信もチェック → URLを控える',
            'YouTubeで最近の登壇動画があれば、そのURLも控える',
        ],
        [
            'Perplexityは出典付きで回答するので、信頼性の高い記事が見つかりやすい',
            '「最近1年以内の記事に絞って」と追加指示すると、より新鮮な情報が得られる',
            'Pro版（月$20）ならDeep Researchも使える',
        ])

    add_step_section(doc, 2,
        'NotebookLMに全情報を投入 → チャットで深掘り',
        'Google NotebookLM（notebooklm.google.com）',
        [
            'NotebookLM（notebooklm.google.com）を開く',
            '「新しいノートブック」を作成する',
            'STEP1で集めたURLを「ソースを追加」で全部入れる',
            'YouTube動画のURLもそのまま入れる（自動文字起こし）',
            'チャット欄で「この方の最近の関心事は？」と質問する',
        ],
        [
            '南場氏は「行きのタクシーの中で」チャットで質問すると語っている',
            '「トランプ政権についてどう考えているか」など具体的な質問が効果的',
            'NotebookLMは無料で使える！Googleアカウントがあればすぐ開始',
        ])

    add_step_section(doc, 3,
        'Circlebackで会議の議事録を自動作成',
        'Circleback（circleback.ai）',
        [
            'Circleback（circleback.ai）でアカウント作成',
            'Google Calendar / Outlook と連携設定する',
            '会議が始まると自動でAIが録音・文字起こし開始',
            '会議終了後、自動で要約・アクションアイテムが生成される',
            '後日「あの会議で〇〇について何と言っていた？」と検索可能',
        ],
        [
            'Zoom / Google Meet / Teams に対応',
            '日本語の会議にも対応している',
            '「誰が何を言ったか」も記録されるので責任の所在が明確になる',
        ])

    add_step_section(doc, 4,
        'Deep Researchで投資・事業判断の調査',
        'ChatGPT Deep Research',
        [
            'ChatGPT（Pro/Plus）を開く',
            'モデル選択で「Deep Research」を選ぶ',
            '「〇〇市場の最新動向と主要プレイヤーを調査して」と入力',
            'AIが数分〜数十分かけて100以上のソースを自動で調査',
            '出典付きの詳細なレポートが生成される',
        ],
        [
            '南場氏は投資判断にこの機能を活用している',
            '従来なら調査チームが数日かかる作業が数十分で完了する',
            'ChatGPT Plus（月$20）以上で利用可能',
        ])

    save_doc(doc, '手順書_南場智子_AI活用術.docx')


# ============================================================
# 2. 堀江貴文
# ============================================================
def create_horie():
    doc = setup_doc()
    add_title_page(doc,
        '堀江 貴文（ホリエモン）のAI活用術',
        'ChatGPTを「もう一人の自分」として使い倒す',
        '「今やらないヤツはバカ！」')

    doc.add_heading('堀江流 AI活用の全体像', level=1)
    doc.add_paragraph(
        '実業家・堀江貴文氏は、ChatGPTの活用法を網羅した書籍『堀江貴文のChatGPT大全』を出版し、'
        '「今やらないヤツはバカ！」と断言するほどのAI推進派です。'
        'スピーチ原稿の生成、書籍の執筆、カスタムGPTの作成、新規事業のアイデア壁打ちなど、'
        '経営者として即座に使える具体的な活用法を実践しています。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        'ChatGPTでスピーチ・文章を一発生成',
        'ChatGPT（chat.openai.com）',
        [
            'ChatGPT（chat.openai.com）を開く',
            '「以下の条件でスピーチ原稿を作って」と指示',
            '条件を具体的に書く（場面・対象者・時間・トーン）',
            '生成された原稿を読み、修正点を追加指示',
            '「もっとカジュアルに」「数字を入れて」等で磨き上げ',
        ],
        [
            '堀江氏は本の帯コメントもAIで生成している',
            '条件を具体的に書くこと（誰に・何分・フォーマルか等）がポイント',
            '最初から完璧を求めず、何度もやり取りして磨くのがコツ',
        ])

    add_step_section(doc, 2,
        'マイGPT（カスタムGPT）で社内専門AIを作る',
        'ChatGPT Plus/Team（GPTs機能）',
        [
            'ChatGPT左メニューの「GPTsを探す」→「作成」をクリック',
            '「あなたは何をするGPTですか？」に役割を入力',
            '「ナレッジ」に自社の資料（PDF/Word等）をアップロード',
            'テスト質問で動作確認 → 修正を繰り返す',
            '完成したら社内メンバーにリンクを共有',
        ],
        [
            '堀江氏は就業規則を読み込ませた「総務部長GPT」を推奨',
            '他にも営業マニュアルGPT、商品知識GPT、クレーム対応GPT等が作れる',
            'ChatGPT Plus（月$20）で利用可能',
        ])

    add_step_section(doc, 3,
        'AIで本を1冊書く（コンテンツ量産）',
        'ChatGPT + 音声入力',
        [
            '自分の考えをスマホの音声メモで録音する',
            'Whisper等の文字起こしツールでテキスト化',
            'ChatGPTに「以下の文字起こしを章立てして書籍にして」と依頼',
            '各章ごとに「もっと具体例を入れて」等で品質を上げる',
            '最終チェックは自分の目で行い、出版・社内共有',
        ],
        [
            '堀江氏は99%をAIが書いたビジネス書を実際に出版した',
            '社長の考えを社内に共有するための「社長の考え方ブック」制作にも応用可能',
            'Whisperは無料で使える文字起こしAI',
        ])

    add_step_section(doc, 4,
        '新規事業のアイデア壁打ち',
        'ChatGPT（GPT-4o / o1）',
        [
            'ChatGPTを開き「新規事業のアイデアを一緒に考えて」と入力',
            '自社の強み・リソース・課題を箇条書きで伝える',
            '出てきたアイデアに「なぜそれが上手くいくと思う？」と深掘り',
            '「リスクと対策を3つずつ挙げて」で現実チェック',
            '良いアイデアは「事業計画書の叩き台を作って」で具体化',
        ],
        [
            'AIは「いい質問をする力」で結果が変わる',
            '漠然と聞くのではなく、具体的な制約条件（予算・期間・リソース等）を与えるのがコツ',
        ])

    save_doc(doc, '手順書_堀江貴文_AI活用術.docx')


# ============================================================
# 3. Sam Altman
# ============================================================
def create_altman():
    doc = setup_doc()
    add_title_page(doc,
        'Sam Altman（OpenAI CEO）のAI活用術',
        '「退屈なタスク」こそAIに任せる',
        '「AIは退屈な仕事にこそ使え」')

    doc.add_heading('Altman流 AI活用の全体像', level=1)
    doc.add_paragraph(
        'OpenAI CEOのSam Altmanは、AIを「退屈なタスク」に使うことを推奨しています。'
        'メールの処理、文書の要約、毎朝のブリーフィングなど、日常的で繰り返しの多い作業にAIを活用し、'
        '自分の時間をより創造的な仕事に充てるアプローチです。'
        'ChatGPT Pulseという機能を「最もお気に入りの機能」と語り、AIが夜間に自動分析した情報を毎朝受け取っています。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        '大量メールの処理をAIで高速化',
        'ChatGPT（chat.openai.com）',
        [
            '処理したいメール本文をコピーする',
            'ChatGPTに「以下のメールを要約して、返信が必要なものをリストアップして」と入力',
            '要約を見て優先順位を判断する',
            '「このメールへの返信を、丁寧に断るトーンで書いて」と依頼',
            '生成された返信を確認 → 微修正してメールに貼り付け',
        ],
        [
            'Altmanは「退屈な仕事にこそAIを使う」と明言している',
            'メール処理は最も効果が出やすい活用法の一つ',
            'まずは1日10通から始めてみよう',
        ])

    add_step_section(doc, 2,
        '長文文書の要約を一瞬で作成',
        'ChatGPT（GPT-4o）',
        [
            '要約したい文書（PDF/Word等）をChatGPTにアップロード',
            '「この文書を3行で要約して」と指示',
            '「重要なポイントを箇条書きで5つ抽出して」と追加指示',
            '「この内容について質問：〇〇はどうなっている？」で深掘り',
            '「経営会議向けの1ページサマリーを作って」で資料化',
        ],
        [
            '契約書、調査レポート、競合資料など読む時間がない文書に最適',
            'PDFは直接アップロードできる（ChatGPT Plus以上）',
            '「専門用語を避けてわかりやすく」と追加指示すると社長向けに最適化される',
        ])

    add_step_section(doc, 3,
        'ChatGPT Pulseで毎朝のブリーフィング',
        'ChatGPT Pulse（最新機能）',
        [
            'ChatGPTアプリを開き、プロフィール設定から「メモリ」をONにする',
            '日頃から関心事をChatGPTに話しかけておく（業界ニュース等）',
            'Pulse機能をONにする（設定 → Pulse）',
            '毎朝、Pulseが「あなた向けの更新情報」を自動生成',
            '関心事の進展、前日の会話のフォローアップ等が届く',
        ],
        [
            'Altmanが「最もお気に入りの機能」と語る注目機能',
            '「秘書が毎朝ブリーフィング」をAIが自動で実現するイメージ',
            'ChatGPT Plus/Proで利用可能',
        ])

    add_step_section(doc, 4,
        '日常生活の「ちょっとした疑問」にAIを使う',
        'ChatGPTアプリ（スマートフォン）',
        [
            'スマホにChatGPTアプリをインストール',
            '音声モードをONにする（マイクアイコンをタップ）',
            '話しかけるだけで回答が返ってくる',
            '「今夜のレシピ提案して」「子供にこう聞かれたら？」等',
            'メモリ機能で好みを覚えてもらい、どんどんパーソナルに',
        ],
        [
            'Altmanは育児の質問にもAIを使っている',
            '音声モードなら運転中や料理中もハンズフリーで使える',
            '「メモリ」機能で家族の好み等を覚えさせると提案の精度が上がる',
        ])

    save_doc(doc, '手順書_SamAltman_AI活用術.docx')


# ============================================================
# 4. Tobi Lutke
# ============================================================
def create_lutke():
    doc = setup_doc()
    add_title_page(doc,
        'Tobi Lutke（Shopify CEO）のAI活用術',
        '「AIで100倍の成果」を出す組織の作り方',
        '「AI活用は全社員の基本的な期待事項」')

    doc.add_heading('Lutke流 AI活用の全体像', level=1)
    doc.add_paragraph(
        'Shopify CEOのTobi Lutkeは、2025年に「AI活用は全社員の基本的な期待事項」と全社メモで宣言。'
        '新規採用の前に「AIで代替できないか？」を証明することを義務づけ、'
        '人事評価にもAI活用度を組み込むなど、最も「制度化」が進んだAI経営を実践しています。'
        'Claude、Cursor、GitHub Copilotなど複数のAIツールを全社員に提供し、'
        '100倍の成果を出した社員の事例も生まれています。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        'アイデアをAIで即プロトタイプ化',
        'Claude（claude.ai）/ ChatGPT',
        [
            '新しいアイデアが浮かんだらすぐにAIを開く',
            '「〇〇のプロトタイプを作りたい。要件は以下：」と入力',
            '要件を箇条書きで列挙（ユーザー・機能・制約）',
            'AIが出したプロトタイプ案に対して「ここをもっと簡単に」等で修正',
            'プロトタイプ案が固まったらチームに共有して議論',
        ],
        [
            'Lutkeは「GSD（Get Shit Done）プロトタイプフェーズ」でのAI活用を重視',
            '完璧なものを作る前に、AIで素早く叩き台を作ることが重要',
            'Claudeの「Artifacts」機能ならHTMLプロトタイプも作れる',
        ])

    add_step_section(doc, 2,
        '「この仕事はAIでできない？」チェック',
        '社内ルールとして導入',
        [
            '新規採用リクエストが来たら、まず「AIで代替できないか？」と問う',
            '具体的に「どのタスクをAIに任せたか？」を報告させる',
            'AIで80%解決できるなら、残り20%を既存メンバーに振り分け',
            'それでも人が必要な場合のみ採用を承認',
            'この判断プロセスを四半期ごとに見直す',
        ],
        [
            'Shopifyでは全チームにこのルールを適用している',
            '経営者から始めることが重要 → 自分が先にAIを使い倒す姿を見せる',
        ])

    add_step_section(doc, 3,
        'AIツールを用途別に使い分ける',
        'Claude / Cursor / GitHub Copilot',
        [
            '長文の分析・戦略立案 → Claude（200Kトークンの長文対応）',
            'コード開発・自動化 → Cursor（AIコードエディタ）',
            '日常のコーディング補助 → GitHub Copilot',
            '情報検索・調査 → Perplexity AI',
            '各ツールの得意分野を把握し、タスクに応じて切り替える',
        ],
        [
            'Claude → 思考のパートナー、Cursor → 開発のパートナー、Copilot → 日常の効率化',
            'まずは1つから始めて、慣れたら用途別に使い分けを覚える',
        ])

    add_step_section(doc, 4,
        '人事評価にAI活用度を組み込む',
        '人事制度の設計',
        [
            'パフォーマンスレビューの評価項目に「AI活用度」を追加',
            '「AIをどのように業務に活用したか」を自己申告させる',
            'チームリーダーがAI活用の好事例を共有する場を設ける',
            'AI活用で成果を上げた社員を表彰する仕組みを作る',
            '四半期ごとにAI活用の成果をチーム単位で振り返る',
        ],
        [
            '「使え」と言うだけでは変わらない → 評価に組み込むことで行動が変わる',
            'Lutkeは「100倍の成果を出した社員」の事例を全社に共有している',
            'まずは小さな成功事例を積み上げることが大事',
        ])

    save_doc(doc, '手順書_TobiLutke_AI活用術.docx')


# ============================================================
# 5. 落合陽一
# ============================================================
def create_ochiai():
    doc = setup_doc()
    add_title_page(doc,
        '落合 陽一（ピクシーダストCEO）のAI活用術',
        'Claude Codeに「寝てる間に作らせる」最先端の使い方',
        '「2026年にはほとんどの知的作業がAIに置き換わる」')

    doc.add_heading('落合流 AI活用の全体像', level=1)
    doc.add_paragraph(
        'メディアアーティスト・研究者であり、ピクシーダストテクノロジーズCEOの落合陽一氏は、'
        'AI活用の最先端を実践するリーダーです。Claude Codeに指示を出して寝ると朝にはツールが完成している'
        'という「寝ている間にAIに作らせる」スタイルや、Notionで全情報をWiki化する手法、'
        'そして「AIの出力から良いものを選ぶ審美眼こそが人間の価値」という考え方は、'
        'あらゆる経営者にとって示唆に富んでいます。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        'Claude Codeに「寝ている間に」タスクを実行させる',
        'Claude Code（ターミナル）',
        [
            'ターミナル（コマンドライン）でClaude Codeを起動',
            '作りたいものを自然言語で詳しく指示する',
            '「完了したらファイルを保存して」と追加指示',
            'AIが自動でコードを書き始める → そのまま就寝',
            '朝起きて結果を確認 → 修正が必要なら追加指示',
        ],
        [
            '落合氏はClaude Codeに「Claude Codeを作って」と指示して寝たら、朝にはツールが完成していた',
            '経営者の場合、Webサイトや社内ツールの作成にも応用可能',
            'プログラミング知識がなくても日本語の指示で動く',
        ])

    add_step_section(doc, 2,
        'Notionで「自分だけのWiki」を構築',
        'Notion（notion.so）',
        [
            'Notion（notion.so）でワークスペースを作成',
            '「社内マニュアル」「議事録」「プロジェクト」等のページを作る',
            'あらゆる情報をNotionに集約する（URLリンクも可）',
            'Notion AI機能で「この文書を要約して」と指示できる',
            'チームメンバーと共有し、全員が同じ情報にアクセス',
        ],
        [
            '落合氏は1人で50人の学生を見ており「Notionがないと仕事が回らない」と語る',
            '社長が持つ暗黙知をNotionに言語化して蓄積すると組織の知識資産になる',
            '無料プランでも十分に使える',
        ])

    add_step_section(doc, 3,
        'AI出力の「審美眼」を鍛える',
        'Claude / ChatGPT / 各種AIツール',
        [
            '同じ質問を複数のAI（ChatGPT, Claude, Gemini）に投げる',
            'それぞれの回答を比較し、最も優れたものを選ぶ',
            '「なぜこの回答が良いと思ったか」を自分で言語化する',
            'AIの出力を「そのまま使わず」必ず自分の視点で編集する',
            'この選別作業を繰り返すことで「審美眼」が磨かれる',
        ],
        [
            '落合氏：AIを上手に使える人と使えない人の差は「審美眼」（選ぶ力）',
            'AIが100個のアイデアを出したら最良の1個を選べるかが勝負',
            '経営判断に直結するスキル',
        ])

    add_step_section(doc, 4,
        'vibe-localでネット不要のAI環境を構築',
        'vibe-local（落合氏開発ツール）',
        [
            'GitHub から vibe-local をダウンロード',
            'ローカルLLM（Ollama等）をインストール',
            'vibe-localを起動すると、ネット不要でAIコーディングが可能',
            '機密情報を含むプロジェクトでも安心して使える',
            '社内ネットワーク外でもAIの力を活用できる',
        ],
        [
            '完全無料・ネットワーク不要で使える',
            '機密情報がクラウドに送信されないためセキュリティ面で安心',
            '外部にデータを出せない業界に最適',
        ])

    save_doc(doc, '手順書_落合陽一_AI活用術.docx')


# ============================================================
# 6. Mustafa Suleyman
# ============================================================
def create_suleyman():
    doc = setup_doc()
    add_title_page(doc,
        'Mustafa Suleyman（Microsoft AI CEO）のAI活用術',
        'Copilotを「AIコンパニオン」として日常に溶け込ませる',
        '「AIに任せられる知的作業は、もうほぼ全てだ」')

    doc.add_heading('Suleyman流 AI活用の全体像', level=1)
    doc.add_paragraph(
        'DeepMind共同創業者であり、現在Microsoft AI CEOを務めるMustafa Suleymanは、'
        'Copilotを「AIコンパニオン」として個人の生活に深く統合しています。'
        '映画のパーソナルデータベース管理から、Copilot for Microsoft 365による業務効率化、'
        'そして「あらゆる知的作業をAIに聞いてみる」習慣の確立まで、'
        'AIを日常の一部にする実践的なアプローチが特徴です。'
    )
    doc.add_paragraph('')

    add_step_section(doc, 1,
        'Copilotを「個人データベース」にする',
        'Microsoft Copilot（copilot.microsoft.com）',
        [
            'Copilot（copilot.microsoft.com）を開く',
            '「メモリ」機能をONにする（設定から）',
            '好きな映画・本・レストラン等をCopilotに伝える',
            '「次に見るべき映画を推薦して」と聞くとパーソナルな提案が返る',
            '仕事の好みや判断基準も伝えておくと業務でも活用できる',
        ],
        [
            'Suleyman自身が映画のパーソナルデータベースとしてCopilotを活用している',
            '経営の意思決定基準を記録させると「先月の判断と矛盾してない？」も聞ける',
        ])

    add_step_section(doc, 2,
        'Copilot for Microsoft 365 で業務を高速化',
        'Microsoft 365 + Copilot',
        [
            'Microsoft 365管理者がCopilotライセンスを有効化',
            'Wordで「このメモから企画書を作って」→ 一瞬で文書生成',
            'Excelで「このデータの傾向を分析して」→ グラフ付きレポート',
            'Teamsで「この会議の要約とアクションアイテムを」→ 自動生成',
            'Outlookで「このメールスレッドを要約して」→ 重要ポイント抽出',
        ],
        [
            '既にWord/Excel/Teamsを使っているなら最も導入ハードルが低いAI',
            'Teamsの会議要約は「参加できなかった会議の内容を5分で把握できる」と好評',
            '月額3,750円/ユーザー（税別）',
        ])

    add_step_section(doc, 3,
        'AIに「自分のことを覚えてもらう」',
        'Copilot / ChatGPT（メモリ機能）',
        [
            'AIとの会話で、自分の好み・判断基準・役割を伝える',
            '「私は〇〇の社長で、△△業界に詳しい」と自己紹介する',
            '仕事の進め方の好みも伝える（データ重視 / 直感重視 等）',
            '過去の会話を踏まえた提案が返ってくるようになる',
            '定期的に「私について何を覚えている？」で確認・修正',
        ],
        [
            'Suleymanのビジョンは「全ての人にAIコンパニオンを」',
            'AIに自分を理解させるほど、提案の質が劇的に向上する',
        ])

    add_step_section(doc, 4,
        '「とりあえずAIに聞いてみる」を習慣化',
        'Copilot / ChatGPT / Claude（全般）',
        [
            '迷ったらまずAIに聞く →「〇〇についてどう思う？」',
            '知らないことがあったらAIに聞く →「〇〇を教えて」',
            '文章を書く前にAIに叩き台を作らせる',
            'スケジュール調整もAIに相談 →「来週の優先順位を整理して」',
            '1日に最低10回はAIに話しかけることを目標にする',
        ],
        [
            'Suleymanは「プロのタスクの大部分は12〜18ヶ月以内にAIで自動化される」と予測',
            '今から「AIに聞く習慣」を身につけておくことが重要',
            '最初は質問の質が低くてもOK → 使い続けることで上手くなる',
        ])

    save_doc(doc, '手順書_MustafaSuleyman_AI活用術.docx')


# ============================================================
if __name__ == "__main__":
    create_nanba()
    create_horie()
    create_altman()
    create_lutke()
    create_ochiai()
    create_suleyman()
    print("\n全6ファイルの作成が完了しました！")
