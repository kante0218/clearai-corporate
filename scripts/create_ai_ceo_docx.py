#!/usr/bin/env python3
"""AI活用で有名な社長6名のWord資料を作成"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== Style setup =====
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Yu Gothic'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ===== CEO Data =====
ceos = [
    {
        "name": "南場 智子（なんば ともこ）",
        "company": "株式会社ディー・エヌ・エー（DeNA）",
        "title": "代表取締役会長",
        "country": "日本",
        "slogan": "「AIにオールイン」宣言",
        "overview": (
            "2025年に「DeNAはAIにオールイン」を宣言し、経営者自ら生成AIのヘビーユーザーとして具体的な活用法を公開。"
            "「7つのAI活用術」として、Perplexity AIでの情報収集、NotebookLMでの資料要約、ChatGPT o1での戦略立案、"
            "Circlebackでの議事録自動作成など、実践的な活用方法を示した。"
        ),
        "ai_usage": [
            "Perplexity AI：初めて会う人の情報を「必読記事は何ですか」と聞いて収集",
            "NotebookLM：収集した記事URLをアップロードし、要点を瞬時に把握",
            "ChatGPT o1：戦略立案やプロトタイピングに活用",
            "Circleback：会議の自動議事録作成",
        ],
        "results": [
            "QA業務の工数を50%削減（従来の半分で同等の成果）",
            "Pococha配信審査コストを60%削減",
            "現業3,000名体制を半分の人員でも発展させる計画を策定",
            "バーティカルAIエージェント事業を新規立ち上げ（スポーツ・ヘルスケア領域）",
        ],
        "takeaway": (
            "経営者自身がAIのヘビーユーザーであることが組織のAI活用を加速させる好例。"
            "具体的なツール名と使い方を公開することで、社員が真似しやすい環境を作っている。"
            "「経営者がAIに興奮しているかがポイント」という言葉は、まさにAI経営の本質を突いている。"
        ),
    },
    {
        "name": "孫 正義（そん まさよし）",
        "company": "ソフトバンクグループ株式会社",
        "title": "代表取締役会長兼社長",
        "country": "日本",
        "slogan": "ASI時代のNo.1プラットフォーマーへ",
        "overview": (
            "AI時代の覇者を目指し、OpenAIに4.8兆円という歴史上最大規模の投資を実行。"
            "「Stargate Project」で世界最大級のAIインフラを構築し、グループ全体で10億のAIエージェントを"
            "作る「Crystal Intelligence」構想を発表。社員1人あたり1,000個のAIエージェントが常時稼働する世界を目指す。"
        ),
        "ai_usage": [
            "Crystal Intelligence：自己増殖型AIエージェント群（世界初の構想）",
            "社員1人あたり1,000個のAIエージェントが24時間365日稼働",
            "ネットワーク機器監視から人事・財務・営業まで全部門にAIを導入",
            "生成AIコンテスト：グループ5万人から累計26万件のAIアイデアを収集",
        ],
        "results": [
            "OpenAIに4.8兆円を投資（歴史上最大規模の未上場企業への投資）",
            "Stargate Projectで米国に世界最大級のAIデータセンターを建設",
            "グループ全体で10億AIエージェント構想を発表",
            "ASI（人工超知能）のNo.1プラットフォーマーを目指すと宣言",
        ],
        "takeaway": (
            "AIへの投資規模とビジョンの大きさは世界でも群を抜いている。"
            "注目すべきは、単なる投資だけでなく、社内で5万人規模の生成AIコンテストを開催し、"
            "ボトムアップでもAI活用を推進している点。トップダウンとボトムアップの両輪で組織を動かしている。"
        ),
    },
    {
        "name": "三木谷 浩史（みきたに ひろし）",
        "company": "楽天グループ株式会社",
        "title": "代表取締役会長兼社長",
        "country": "日本",
        "slogan": "「No AI, No Future」宣言",
        "overview": (
            "「AIなくして未来なし」を掲げ、楽天市場を世界最先端のAI活用ECプラットフォームへ進化させると宣言。"
            "「AIの民主化」をキーワードに、全従業員3万人がAIを活用できる環境を整備。"
            "2030年に楽天市場の流通総額10兆円達成を目指す。"
        ),
        "ai_usage": [
            "Rakuten AI：独自開発の生成AIを楽天市場に導入（2025年秋〜）",
            "OpenAIとの戦略的パートナーシップ",
            "「トリプル20」：マーケティング・オペレーション・顧客効率をそれぞれ20%向上",
            "全従業員3万人がAI活用、毎日の利用者は8,000人を突破",
        ],
        "results": [
            "マーケティング効率20%向上を達成",
            "オペレーション効率の大幅改善",
            "2030年に楽天市場の流通総額10兆円へ向けた成長戦略を策定",
            "「AIの民主化」で全社員のデジタルスキル底上げ",
        ],
        "takeaway": (
            "「AIの民主化」というコンセプトが秀逸。一部の技術者だけでなく、全社員がAIを使えるようにすることで、"
            "組織全体の生産性を底上げしている。EC事業者にとって、楽天のAI活用はそのまま参考にできる事例が多い。"
            "「AIを使いこなせない企業・人の未来は厳しい」という警鐘は、経営者として重く受け止めるべきメッセージ。"
        ),
    },
    {
        "name": "Satya Nadella（サティア・ナデラ）",
        "company": "Microsoft Corporation",
        "title": "CEO（最高経営責任者）",
        "country": "アメリカ",
        "slogan": "AIを「認知増幅装置」として全製品に統合",
        "overview": (
            "Copilotブランドの下、Office、Azure、GitHub等あらゆるMicrosoft製品にAIを統合。"
            "OpenAIへの巨額投資を主導し、エンタープライズAIの世界標準を構築した。"
            "AIを「認知増幅装置（Cognitive Amplifier）」と定義し、人間の能力を置き換えるのではなく増幅するものと位置づけている。"
        ),
        "ai_usage": [
            "Microsoft Copilot：Word、Excel、Teams等のOffice全製品にAIアシスタントを統合",
            "Azure OpenAI Service：企業が自社データでAIを活用できるプラットフォーム",
            "GitHub Copilot：プログラマーの生産性を劇的に向上させるAIペアプログラマー",
            "Copilot+ PC：AIネイティブなハードウェアの展開",
        ],
        "results": [
            "AI事業が年間売上130億ドル超を達成",
            "Azure成長率がAI需要で大幅に加速",
            "GitHub Copilotが開発者のコーディング時間を55%短縮",
            "2026年を「AIの転換年（Pivotal Year）」と位置づけ",
        ],
        "takeaway": (
            "Nadellaの最大の功績は、AIを「特別なもの」ではなく「日常のツール」にしたこと。"
            "WordやExcelといった誰もが使うツールにAIを統合することで、技術リテラシーに関わらず全ての人がAIの恩恵を受けられる。"
            "「AIは人間を置き換えるのではなく増幅する」という哲学は、AI導入に不安を感じる社員への最良のメッセージになる。"
        ),
    },
    {
        "name": "Marc Benioff（マーク・ベニオフ）",
        "company": "Salesforce, Inc.",
        "title": "会長兼CEO",
        "country": "アメリカ",
        "slogan": "社内業務の50%をAIが担当",
        "overview": (
            "Agentforceプラットフォームで「デジタル従業員」という新概念を推進。"
            "社内業務の30〜50%をAIが処理する体制を実現し、顧客対応の85%をAIが自動解決。"
            "ダボス会議で「我々は全員が人間だけの組織を率いる最後のCEO世代だ」と発言し、大きな話題を呼んだ。"
        ),
        "ai_usage": [
            "Agentforce：AIエージェントプラットフォーム（会社名変更を検討するほど注力）",
            "顧客対応の85%をAIエージェントが自動解決",
            "営業リードの選別をAIで40%高速化",
            "エンジニアリング、カスタマーサービス、マーケティング全般をAIが支援",
        ],
        "results": [
            "社内業務の30〜50%をAIが処理する体制を実現",
            "カスタマーサービス部門で約4,000名分の業務をAIが代替",
            "2030年に年間売上600億ドル達成を目標（Agentforce効果）",
            "採用の51%が社内リデプロイメント（外部採用抑制）",
        ],
        "takeaway": (
            "Benioffの事例は「AI導入の最も進んだ姿」を示している。業務の50%をAIが担当するというのは驚異的な数字。"
            "重要なのは、単に人を減らすのではなく、社内の人材をAIでは担えない業務に再配置している点。"
            "「人間にしかできない価値」を明確にした上でAIを導入することが、組織の持続的な成長につながる。"
        ),
    },
    {
        "name": "Tobi Lutke（トビ・リュトケ）",
        "company": "Shopify Inc.",
        "title": "CEO（最高経営責任者）",
        "country": "カナダ",
        "slogan": "「AI活用は全社員の基本的な期待事項」",
        "overview": (
            "2025年に全社員向けメモで「AI活用は全社員の基本的な期待事項」と宣言。"
            "新規採用の前に「この仕事はAIで代替できないのか？」を証明することを義務づけた。"
            "人事評価にもAI活用度を組み込み、組織全体をAI-firstに転換した。"
        ),
        "ai_usage": [
            "全社員にAI活用を「基本的な期待事項」として義務化",
            "新規採用前に「AI代替不可」の証明を義務づけ",
            "人事評価にAI活用度を反映",
            "Shopify Magic：EC事業者向けAIツール群（商品説明自動生成等）",
        ],
        "results": [
            "組織全体にAI-first文化を確立",
            "不要な新規採用を抑制し、コスト最適化を実現",
            "EC事業者にAIツール「Shopify Magic」を提供し競合との差別化",
            "AI活用度を人事評価の重要指標として組み込み",
        ],
        "takeaway": (
            "Lutkeのアプローチは最も「制度化」が進んでいる。「AIを使え」と言うだけでなく、"
            "採用基準と人事評価にAI活用を組み込むことで、組織全体の行動を変えている。"
            "特に「新規採用前にAI代替を検討する」というルールは、あらゆる企業が即座に導入できる施策として参考になる。"
        ),
    },
]

# ===== Title Page =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI活用で世界をリードする経営者たち')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('〜 6人のトップリーダーに学ぶ、AI経営の最前線 〜')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n2025-2026年 最新動向まとめ')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\nclearAI')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

doc.add_page_break()

# ===== Table of Contents =====
doc.add_heading('目次', level=1)
doc.add_paragraph('')

for i, ceo in enumerate(ceos, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {ceo["name"]}（{ceo["company"]}）')
    run.font.size = Pt(13)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'    {ceo["slogan"]}')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.font.name = 'Yu Gothic'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

p = doc.add_paragraph('')
run = p.add_run('7. AI経営の共通ポイント（まとめ）')
run.font.size = Pt(13)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

doc.add_page_break()

# ===== Individual CEO Pages =====
for i, ceo in enumerate(ceos, 1):
    doc.add_heading(f'{i}. {ceo["name"]}', level=1)

    # Basic info table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cells_data = [
        ('会社名', ceo['company']),
        ('役職', ceo['title']),
        ('国', ceo['country']),
        ('キーワード', ceo['slogan']),
    ]
    for row_idx, (label, value) in enumerate(cells_data):
        cell_label = table.cell(row_idx, 0)
        cell_label.text = label
        for p in cell_label.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        cell_value = table.cell(row_idx, 1)
        cell_value.text = value

    doc.add_paragraph('')

    # Overview
    doc.add_heading('概要', level=2)
    p = doc.add_paragraph(ceo['overview'])
    p.paragraph_format.space_after = Pt(6)

    # AI Usage
    doc.add_heading('AI活用の具体的な方法', level=2)
    for item in ceo['ai_usage']:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    # Results
    doc.add_heading('主な成果・インパクト', level=2)
    for item in ceo['results']:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    # Takeaway
    doc.add_heading('経営者への示唆', level=2)
    p = doc.add_paragraph(ceo['takeaway'])
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

# ===== Summary Section =====
doc.add_heading('7. AI経営の共通ポイント（まとめ）', level=1)
doc.add_paragraph('')

summary_points = [
    (
        '経営者自身がAIのヘビーユーザーであること',
        '南場氏のように経営者自らAIを使い倒し、具体的な活用法を社員に示すことで組織全体のAI活用が加速します。'
        '「使え」と言うだけでなく「こう使う」を見せることが重要です。'
    ),
    (
        'AI前提で組織・業務を再設計すること',
        'Benioff氏やLutke氏のように、人事制度・採用基準・業務フローをAI前提で再構築することが成功の鍵です。'
        '既存の業務にAIを「足す」のではなく、AI前提で「作り直す」発想が必要です。'
    ),
    (
        '全社員をAI人材化する仕組みを作ること',
        '三木谷氏の「AIの民主化」、孫氏の「生成AIコンテスト」のように、全社員がAIを使える環境と動機を整備することが'
        '組織全体の底上げにつながります。'
    ),
    (
        'AIエージェントが次のフロンティア',
        '6名全員がAIエージェント（自律的に業務を遂行するAI）を次の主戦場と位置づけています。'
        '単なるチャットボットではなく、業務を自律的に遂行するAIの導入が次のステップです。'
    ),
]

for i, (title, desc) in enumerate(summary_points, 1):
    doc.add_heading(f'ポイント{i}: {title}', level=2)
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(12)

# ===== Closing =====
doc.add_page_break()
doc.add_heading('おわりに', level=1)
p = doc.add_paragraph(
    '本資料で紹介した6名のリーダーに共通するのは、「AIは未来の話ではなく、今すぐ取り組むべき経営課題」'
    'という強い危機感と行動力です。'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    '彼らのアプローチには規模の違いはあれど、中小企業でも参考にできるエッセンスが詰まっています。'
    'まずは経営者自身がAIツールを使ってみること。そして、社員が使える環境と仕組みを整えること。'
    'この2つから始めることが、AI経営への第一歩です。'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    'AIの導入・活用についてお悩みの際は、ぜひclearAIにご相談ください。'
    '御社の状況に合わせた最適なAI活用プランをご提案いたします。'
)
p.paragraph_format.space_after = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('clearAI')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x96, 0xFF)

# Save
output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "AI活用リーダー6名.docx")
doc.save(output_path)
print(f"Word document saved: {output_path}")
